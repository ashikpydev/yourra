"""
Word (.docx) generator for a finished, ready-to-hand-over transcript.

Produces a polished document:
  - Title
  - Demographic / metadata table (survey type + respondent details) — the
    standard header used for qualitative research deliverables
  - Bangla verbatim transcript (speaker-labelled, Bangla font)
  - English translation (speaker-labelled)
"""
import io
import json
import re
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

BANGLA_FONT = "Nirmala UI"  # ships with Windows; renders Bangla cleanly
BRAND = RGBColor(0x4F, 0x46, 0xE5)

# Demographic fields shown (in order). Always rendered so the table layout is
# consistent, even when a value wasn't provided.
META_FIELDS = [
    ("survey_type", "Survey type"),
    ("resp_name", "Respondent name / ID"),
    ("resp_age", "Age"),
    ("resp_sex", "Sex"),
    ("resp_education", "Education"),
    ("resp_profession", "Profession"),
    ("resp_location", "Location"),
    ("interviewer", "Interviewer / Moderator"),
    ("interview_date", "Interview date"),
]


def _bangla(run):
    run.font.name = BANGLA_FONT
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), BANGLA_FONT)
    rfonts.set(qn("w:hAnsi"), BANGLA_FONT)
    rfonts.set(qn("w:cs"), BANGLA_FONT)


_TS_RE = re.compile(r"^\[\d{1,2}:\d{2}(?::\d{2})?\]$")


def _add_transcript(doc, text, bangla=False):
    for raw in (text or "").split("\n"):
        line = raw.rstrip()
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        if not line:
            continue
        stripped = line.strip()
        # Timestamp line, e.g. [0:02:00]
        if _TS_RE.match(stripped):
            r = p.add_run(stripped)
            r.bold = True
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
            continue
        # Speaker line: bold the label before the first colon
        if ":" in line and not stripped.startswith("[") and len(line.split(":", 1)[0]) <= 24:
            speaker, rest = line.split(":", 1)
            r1 = p.add_run(speaker + ":")
            r1.bold = True
            r2 = p.add_run(rest)
            if bangla:
                _bangla(r1); _bangla(r2)
        else:
            r = p.add_run(line)
            if bangla:
                _bangla(r)


def build_transcript_docx(job: dict) -> bytes:
    meta = {}
    raw = job.get("respondent_meta")
    if raw:
        try:
            meta = json.loads(raw)
        except Exception:
            meta = {}

    doc = Document()

    # Title
    title = doc.add_heading("Qualitative Research Transcript", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sub = doc.add_paragraph()
    sr = sub.add_run("Prepared with YourRA — Bangla verbatim + English translation")
    sr.italic = True
    sr.font.size = Pt(10)
    sr.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    # Demographic table
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    table.columns[0].width = Pt(150)

    def add_row(label, value):
        cells = table.add_row().cells
        lr = cells[0].paragraphs[0].add_run(label)
        lr.bold = True
        cells[1].text = "" if value in (None, "") else str(value)

    for key, label in META_FIELDS:
        add_row(label, meta.get(key, ""))
    # System fields
    add_row("Audio file", job.get("original_filename", ""))
    dur = job.get("duration_minutes")
    add_row("Audio length", f"{dur} min" if dur else "")
    add_row("Engine", "Google Gemini 2.5 " + ((job.get("model_used") or "").capitalize() or "Pro"))
    add_row("Generated", datetime.now().strftime("%d %b %Y"))

    doc.add_paragraph()

    # Bangla
    doc.add_heading("Bangla Verbatim", level=1)
    _add_transcript(doc, job.get("transcript_bn"), bangla=True)

    # English
    doc.add_heading("English Translation", level=1)
    _add_transcript(doc, job.get("transcript_en"), bangla=False)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()
